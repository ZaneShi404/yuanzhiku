"""Local parsing adapters with an explicit, offline-only Docling policy."""

from __future__ import annotations

import hashlib
import io
import json
from pathlib import Path
from typing import Any

from app.domain.parsing import ParsedDocument, ParsedSegment


def _config_hash(name: str, version: str) -> str:
    return hashlib.sha256(f"{name}:{version}:local-only".encode("ascii")).hexdigest()


def _read_limited(artifact_path: Path, maximum_bytes: int | None = None) -> bytes:
    size = artifact_path.stat().st_size
    if maximum_bytes is not None and size > maximum_bytes:
        raise ValueError("解析输入超过进程内存限制")
    with artifact_path.open("rb") as stream:
        return stream.read()


def parse_local(artifact_path: Path, filename: str, media_type: str | None = None, maximum_bytes: int | None = None) -> ParsedDocument:
    """Parse with the existing fully local fallback adapters only."""
    suffix = Path(filename).suffix.lower()
    raw = _read_limited(artifact_path, maximum_bytes)
    if suffix in {".txt", ".md", ".markdown"}:
        try:
            return ParsedDocument(raw.decode("utf-8"), "native-utf8", _config_hash("native-utf8", "1"), suffix[1:])
        except UnicodeDecodeError:
            return ParsedDocument("", "native-utf8", _config_hash("native-utf8", "1"), suffix[1:], "文本不是有效 UTF-8")
    if suffix == ".pdf" or media_type == "application/pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            if reader.is_encrypted:
                return ParsedDocument("", "pypdf-local", _config_hash("pypdf-local", "5.4.0"), "pdf", "PDF 已加密，等待人工提供可解析副本")
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(pages)
            if not text.strip():
                return ParsedDocument("", "pypdf-local", _config_hash("pypdf-local", "5.4.0"), "pdf", "awaiting_ocr")
            segments: list[ParsedSegment] = []
            offset = 0
            for page_number, page_text in enumerate(pages, start=1):
                if page_text:
                    segments.append(ParsedSegment(offset, offset + len(page_text), {
                        "type": "pdf_page_char_range", "page": page_number, "char_range": [0, len(page_text)],
                    }))
                offset += len(page_text) + 2
            return ParsedDocument(text, "pypdf-local", _config_hash("pypdf-local", "5.4.0"), "pdf", segments=tuple(segments))
        except Exception:
            return ParsedDocument("", "pypdf-local", _config_hash("pypdf-local", "5.4.0"), "pdf", "PDF 无法本地解析")
    if suffix == ".docx" or media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document

            document = Document(io.BytesIO(raw))
            paragraphs = [(ordinal, paragraph.text) for ordinal, paragraph in enumerate(document.paragraphs, start=1) if paragraph.text]
            text = "\n\n".join(paragraph for _, paragraph in paragraphs)
            if not text.strip():
                return ParsedDocument("", "python-docx-local", _config_hash("python-docx-local", "1.1.2"), "docx", "DOCX 没有可提取文本")
            segments: list[ParsedSegment] = []
            offset = 0
            for ordinal, paragraph in paragraphs:
                segments.append(ParsedSegment(offset, offset + len(paragraph), {
                    "type": "docx_structure_char_range", "structure": "body", "paragraph_ordinal": ordinal, "char_range": [0, len(paragraph)],
                }))
                offset += len(paragraph) + 2
            return ParsedDocument(text, "python-docx-local", _config_hash("python-docx-local", "1.1.2"), "docx", segments=tuple(segments))
        except Exception:
            return ParsedDocument("", "python-docx-local", _config_hash("python-docx-local", "1.1.2"), "docx", "DOCX 无法本地解析")
    return ParsedDocument("", "unsupported-local", _config_hash("unsupported-local", "1"), suffix.lstrip("."), "不支持的文档类型")


def _docling_segments(document: Any) -> tuple[str, list[ParsedSegment]] | None:
    """按 Docling 条目 provenance 页码聚合文本，产出 pdf_page_char_range segments。

    REQ-021：PDF 证据 locator 必须含真实页码。任何非空文本条目缺少页码时
    返回 None（调用方回退整文路径），绝不产出 page 未知的页级证据。
    """
    pages: dict[int, list[str]] = {}
    seen_items = False
    for item in getattr(document, "texts", []) or []:
        item_text = getattr(item, "text", "") or ""
        if not item_text.strip():
            continue
        seen_items = True
        prov = list(getattr(item, "prov", None) or [])
        page_no = getattr(prov[0], "page_no", None) if prov else None
        if not isinstance(page_no, int) or page_no < 1:
            return None
        pages.setdefault(page_no, []).append(item_text)
    if not seen_items:
        return None
    segments: list[ParsedSegment] = []
    page_texts: list[str] = []
    offset = 0
    for page_no in sorted(pages):
        page_text = "\n".join(pages[page_no])
        page_texts.append(page_text)
        segments.append(ParsedSegment(offset, offset + len(page_text), {
            "type": "pdf_page_char_range", "page": page_no, "char_range": [0, len(page_text)],
        }))
        offset += len(page_text) + 2
    return "\n\n".join(page_texts), segments


class LocalDocumentParser:
    """Docling-first only when an approved, complete local cache is present.

    The current model lock intentionally authorizes no model downloads. This
    adapter never triggers Docling's implicit first-use network downloads.
    """

    def __init__(self, models_directory: Path, lockfile: Path) -> None:
        self.models_directory = models_directory
        self.lockfile = lockfile

    def _model_status(self) -> tuple[bool, str, dict[str, Any]]:
        try:
            lock = json.loads(self.lockfile.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return False, "模型锁文件无效", {}
        models = lock.get("models")
        if not isinstance(models, list) or not models:
            return False, "没有已批准的 Docling 模型", {}
        required_fields = ("name", "version", "source_url", "license", "cache_path", "sha256")
        for model in models:
            if not isinstance(model, dict):
                return False, "模型锁文件无效", {}
            # REQ-013：预批准条目必须锁定版本/来源/许可证/哈希，缺一不予使用。
            if any(not isinstance(model.get(field), str) or not model.get(field) for field in required_fields):
                return False, "模型锁文件缺少必需字段（name/version/source_url/license/cache_path/sha256）", {}
            relative_path = model["cache_path"]
            expected_hash = model["sha256"]
            candidate = self.models_directory / relative_path
            if not candidate.is_file():
                return False, "已批准的 Docling 模型未缓存", {}
            digest = hashlib.sha256(candidate.read_bytes()).hexdigest()
            if digest != expected_hash:
                return False, "已批准的 Docling 模型哈希不匹配", {}
        return True, "", lock

    def capability(self) -> dict[str, object]:
        ready, reason, lock = self._model_status()
        try:
            import docling  # type: ignore[import-not-found]

            package_available = True
            package_version = getattr(docling, "__version__", "unknown")
        except ImportError:
            package_available = False
            package_version = None
        configured = len(lock.get("models", [])) if lock else 0
        enabled = ready and package_available
        return {
            "preferred": "docling",
            "enabled": enabled,
            "configured_model_downloads": configured,
            "package_version": package_version,
            "unavailable_reason": None if enabled else (reason or "Docling Python 包未安装"),
            "fallbacks": ["pypdf", "python-docx", "native-utf8"],
            "cloud_fallback": False,
        }

    def parse(self, artifact_path: Path, filename: str, media_type: str | None, workspace: Path, maximum_bytes: int | None = None) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        # TXT/Markdown are deliberately native UTF-8. PDF uses Docling only
        # after the cache is completely verified; otherwise no implicit network
        # request is permitted and the safe local fallback is selected.
        ready, _, _ = self._model_status()
        if suffix == ".pdf" and ready:
            try:
                from docling.document_converter import DocumentConverter  # type: ignore[import-not-found]

                workspace.mkdir(parents=True, exist_ok=True)
                result = DocumentConverter().convert(str(artifact_path))
                extracted = _docling_segments(result.document)
                if extracted is not None:
                    text, segments = extracted
                else:
                    # 条目缺页码 provenance：整文路径，evidence 走原生兜底 locator。
                    text, segments = result.document.export_to_markdown(), []
                if text.strip():
                    return ParsedDocument(
                        text, "docling-local", _config_hash("docling-local", "approved-cache"), "pdf",
                        segments=tuple(segments),
                    )
            except Exception:
                # No exception content reaches durable job messages. The native
                # fallback is preferred over a cloud or first-use model download.
                pass
        return parse_local(artifact_path, filename, media_type, maximum_bytes)
