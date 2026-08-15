from __future__ import annotations

import io
import json
import os
import queue
import shutil
import sqlite3
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

import pytest
from fastapi.testclient import TestClient

from app.core.config import InstanceLock, choose_port, data_paths
from app.domain.models import PasteImportRequest
from app.main import create_app
from app.services.jobs import JobService, ParserCircuitBreaker


RUN_ROOT = Path(os.environ.get("YUANZHIKU_TEST_RUNTIME", Path(__file__).resolve().parents[1] / "runtime")) / "defect-fixes"


@pytest.fixture()
def runtime_root() -> Path:
    RUN_ROOT.mkdir(parents=True, exist_ok=True)
    root = RUN_ROOT / uuid.uuid4().hex
    root.mkdir()
    yield root
    shutil.rmtree(root, ignore_errors=True)


@pytest.fixture()
def client(runtime_root: Path):
    app = create_app(runtime_root, acquire_lock=False)
    with TestClient(app) as test_client:
        yield test_client


def paste(client: TestClient, title: str, text: str = "# source\n\nshared text") -> dict:
    response = client.post("/api/v1/imports/paste", json={
        "title": title,
        "text": text,
        "rights": "owned",
        "categories": ["document"],
        "tags": ["devfix"],
    })
    assert response.status_code == 201, response.text
    return response.json()


def parse(client: TestClient) -> dict:
    response = client.post("/api/v1/jobs/run-once")
    assert response.status_code == 200, response.text
    return response.json()["job"]


def artifact_files(root: Path) -> list[Path]:
    return [path for path in (root / "artifacts").rglob("*") if path.is_file()]


def test_port_preference_persists_explicit_selection_and_never_rewrites_busy_value(
    runtime_root: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    paths = data_paths(runtime_root / "port-preference")
    available = {18080, 18081}
    monkeypatch.setattr("app.core.config._available", lambda port: port in available)

    assert choose_port(paths, 18080) == 18080
    assert json.loads(paths.port_file.read_text(encoding="utf-8")) == {"port": 18080}
    assert choose_port(paths) == 18080

    assert choose_port(paths, 18081) == 18081
    assert json.loads(paths.port_file.read_text(encoding="utf-8")) == {"port": 18081}

    with pytest.raises(RuntimeError, match="未修改保存的端口偏好"):
        choose_port(paths, 18082)
    assert json.loads(paths.port_file.read_text(encoding="utf-8")) == {"port": 18081}

    available.remove(18081)
    with pytest.raises(RuntimeError, match="未修改保存的端口偏好"):
        choose_port(paths)
    assert json.loads(paths.port_file.read_text(encoding="utf-8")) == {"port": 18081}


def test_instance_lock_acquisition_never_grows_lock_file(runtime_root: Path) -> None:
    lock_path = runtime_root / "locks" / "instance.lock"
    for _ in range(5):
        lock = InstanceLock(lock_path)
        lock.acquire()
        try:
            assert lock_path.stat().st_size == 1
        finally:
            lock.release()
    assert lock_path.read_bytes() == b"0"


def test_ingest_failure_compensates_only_new_artifact(runtime_root: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services
    original = services.repository.create_ingest

    def failed_ingest(*args, **kwargs):
        raise sqlite3.IntegrityError("injected")

    monkeypatch.setattr(services.repository, "create_ingest", failed_ingest)
    with pytest.raises(sqlite3.IntegrityError):
        services.imports.paste(PasteImportRequest(title="failure", text="new artifact", rights="owned"))
    assert artifact_files(runtime_root) == []
    monkeypatch.setattr(services.repository, "create_ingest", original)

    retained = services.imports.paste(PasteImportRequest(title="dedup", text="shared", rights="owned"))
    monkeypatch.setattr(services.repository, "create_ingest", failed_ingest)
    with pytest.raises(sqlite3.IntegrityError):
        services.imports.paste(PasteImportRequest(title="dedup failure", text="shared", rights="owned"))
    assert services.artifacts.verify(retained["artifact"]["sha256"])
    assert len(artifact_files(runtime_root)) == 1


def test_concurrent_duplicate_ingest_shares_artifact_without_orphans(runtime_root: Path) -> None:
    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services

    def ingest(index: int) -> dict:
        return services.imports.paste(PasteImportRequest(title=f"duplicate-{index}", text="same concurrent bytes", rights="owned"))

    with ThreadPoolExecutor(max_workers=2) as executor:
        results = list(executor.map(ingest, range(2)))
    assert len({result["source"]["id"] for result in results}) == 2
    assert len({result["artifact"]["sha256"] for result in results}) == 1
    assert len(artifact_files(runtime_root)) == 1
    assert len(services.repository.rows_for_export()["sources"]) == 2
    assert len(services.repository.rows_for_export()["content_versions"]) == 2


def test_backups_have_unique_archives_and_retention_never_advertises_missing(client: TestClient, runtime_root: Path) -> None:
    from datetime import UTC, datetime

    with patch("app.services.transfers.datetime") as clock:
        clock.now.return_value = datetime(2030, 1, 2, 3, 4, 5, tzinfo=UTC)
        first = client.post("/api/v1/backups")
        second = client.post("/api/v1/backups")
    assert first.status_code == second.status_code == 201
    assert first.json()["archive_name"] != second.json()["archive_name"]
    records = client.get("/api/v1/backups").json()
    retained_names = {record["archive_name"] for record in records}
    assert second.json()["archive_name"] in retained_names
    assert len(retained_names) == len(records) == 1
    assert all((runtime_root / "backups" / record["archive_name"]).is_file() for record in records if record["state"] == "succeeded")


def build_two_page_pdf() -> bytes:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R 4 0 R] /Count 2 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 6 0 R >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 7 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length 43 >>\nstream\nBT /F1 12 Tf 72 720 Td (first page evidence) Tj ET\nendstream",
        b"<< /Length 44 >>\nstream\nBT /F1 12 Tf 72 720 Td (second page evidence) Tj ET\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for ordinal, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{ordinal} 0 obj\n".encode())
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    output.extend(b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:]))
    output.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
    return bytes(output)


def test_pdf_evidence_uses_each_page_ordinal(client: TestClient) -> None:
    upload = client.post(
        "/api/v1/imports/file",
        data={"rights": "owned", "title": "multi page", "categories": "[]", "tags": "[]", "language": "zh"},
        files={"file": ("two.pdf", build_two_page_pdf(), "application/pdf")},
    )
    assert upload.status_code == 201, upload.text
    parse(client)
    representation = client.get(f"/api/v1/documents/{upload.json()['content_version']['id']}/representations").json()[0]
    evidence = client.get(f"/api/v1/representations/{representation['id']}/evidence").json()
    assert [item["locator"]["page"] for item in evidence] == [1, 2]
    assert [item["excerpt"] for item in evidence] == ["first page evidence", "second page evidence"]

def test_docx_evidence_uses_each_paragraph_ordinal(client: TestClient) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("first paragraph evidence")
    document.add_paragraph("second paragraph evidence")
    output = io.BytesIO()
    document.save(output)
    upload = client.post(
        "/api/v1/imports/file",
        data={"rights": "owned", "title": "multi paragraph", "categories": "[]", "tags": "[]", "language": "zh"},
        files={"file": ("two.docx", output.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert upload.status_code == 201, upload.text
    parse(client)
    representation = client.get(f"/api/v1/documents/{upload.json()['content_version']['id']}/representations").json()[0]
    evidence = client.get(f"/api/v1/representations/{representation['id']}/evidence").json()
    assert [item["locator"]["paragraph_ordinal"] for item in evidence] == [1, 2]
    assert [item["excerpt"] for item in evidence] == ["first paragraph evidence", "second paragraph evidence"]


def test_reimport_unique_card_conflict_returns_409_before_artifact_copy(client: TestClient, runtime_root: Path) -> None:
    url = "https://example.test/existing"
    existing = client.post("/api/v1/external/cards", json={"url": url, "title": "target existing"})
    assert existing.status_code == 201
    donor_root = runtime_root / "donor"
    donor_app = create_app(donor_root, acquire_lock=False)
    donor = donor_app.state.services
    donor.external_cards.create(type("Card", (), {"url": url, "title": "donor existing", "author": None, "notes": None, "tags": []})())
    donor.imports.paste(PasteImportRequest(title="donor artifact", text="artifact that must not be copied", rights="owned"))
    exported = donor.transfers.create_export(True)
    before = artifact_files(runtime_root)
    response = client.post("/api/v1/reimports", json={"archive_path": exported["archive_path"]})
    assert response.status_code == 409, response.text
    assert response.json()["detail"]["code"] == "reimport_conflict"
    assert response.json()["detail"]["message"] == "导入逻辑记录冲突"
    assert response.json()["detail"]["reason"] == "逻辑链或唯一约束冲突，已拒绝且未写入"
    assert response.json()["detail"]["conflicts"]
    assert artifact_files(runtime_root) == before


def test_purge_shared_soft_deleted_artifact_waits_for_last_reference(client: TestClient, runtime_root: Path) -> None:
    first = paste(client, "one")
    second = paste(client, "two")
    parse(client)
    parse(client)
    assert first["artifact"]["sha256"] == second["artifact"]["sha256"]
    client.post(f"/api/v1/sources/{first['source']['id']}/delete")
    client.post(f"/api/v1/sources/{second['source']['id']}/delete")
    purged_first = client.post(f"/api/v1/sources/{first['source']['id']}/purge")
    assert purged_first.status_code == 200, purged_first.text
    assert (runtime_root / "artifacts" / first["artifact"]["sha256"][:2] / first["artifact"]["sha256"]).is_file()
    purged_second = client.post(f"/api/v1/sources/{second['source']['id']}/purge")
    assert purged_second.status_code == 200, purged_second.text
    assert not artifact_files(runtime_root)


def test_parser_circuit_breaker_and_configured_retry_are_observable(runtime_root: Path) -> None:
    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services
    services.repository.update_settings({"max_retry_attempts": 0, "parser_timeout_seconds": 60, "parser_no_progress_seconds": 60})
    imported = services.imports.paste(PasteImportRequest(title="timeout", text="timeout text", rights="owned"))
    assert imported["job"]["max_attempts"] == 0

    def timeout_runner(path, filename, media_type, timeout, no_progress, cancelled, heartbeat):
        assert timeout == 60
        assert no_progress == 60
        heartbeat()
        raise ParserCircuitBreaker("解析无进展断路器已触发")

    services.jobs = JobService(services.repository, services.artifacts, services.documents, services.transfers.create_backup, parse_runner=timeout_runner)
    job = services.jobs.run_once()
    assert job["state"] == "failed"
    assert job["message"] == "解析无进展断路器已触发"
    assert services.repository.get_source(imported["source"]["id"])["processing_state"] == "failed"


def test_permanent_delete_audit_retention_preserves_current_and_other_events(runtime_root: Path) -> None:
    from datetime import UTC, datetime, timedelta

    services = create_app(runtime_root, acquire_lock=False).state.services
    cutoff = datetime.now(UTC) - timedelta(days=366)
    with services.repository.connection() as connection:
        connection.execute(
            "INSERT INTO audit_events(id,event_type,entity_id,result,created_at) VALUES(?,?,?,?,?)",
            ("expired", "source_permanent_delete", "expired-source", "succeeded", (cutoff - timedelta(microseconds=1)).isoformat()),
        )
        connection.execute(
            "INSERT INTO audit_events(id,event_type,entity_id,result,created_at) VALUES(?,?,?,?,?)",
            ("retained", "source_permanent_delete", "retained-source", "succeeded", (cutoff + timedelta(days=1)).isoformat()),
        )
        connection.execute(
            "INSERT INTO audit_events(id,event_type,entity_id,result,created_at) VALUES(?,?,?,?,?)",
            ("other", "source_soft_delete", "other-source", "succeeded", (cutoff - timedelta(days=1)).isoformat()),
        )

    removed = services.repository.prune_source_permanent_delete_audit_events()

    with services.repository.connection() as connection:
        ids = {row["id"] for row in connection.execute("SELECT id FROM audit_events").fetchall()}
    assert removed == 1
    assert ids == {"retained", "other"}
    with pytest.raises(ValueError, match="不得少于 366 天"):
        services.repository.prune_source_permanent_delete_audit_events(365)


def test_permanent_delete_audit_precedes_cleanup_and_rejects_invalid_source(runtime_root: Path) -> None:
    services = create_app(runtime_root, acquire_lock=False).state.services
    with pytest.raises(KeyError, match="来源不存在"):
        services.lifecycle.purge("missing")
    with services.repository.connection() as connection:
        assert connection.execute("SELECT COUNT(*) AS n FROM audit_events").fetchone()["n"] == 0

    imported = services.imports.paste(PasteImportRequest(title="to purge", text="purge text", rights="owned"))
    source_id = imported["source"]["id"]
    services.lifecycle.delete(source_id)
    services.lifecycle.purge(source_id)

    with services.repository.connection() as connection:
        rows = connection.execute(
            "SELECT result FROM audit_events WHERE event_type=? AND entity_id=? ORDER BY created_at ASC",
            ("source_permanent_delete", source_id),
        ).fetchall()
    assert [row["result"] for row in rows] == ["started", "succeeded"]


    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services
    observed: list[int] = []
    worker = JobService(
        services.repository,
        services.artifacts,
        services.documents,
        integrity_runner=lambda sample_size: observed.append(sample_size) or {"valid": True, "checked": sample_size},
    )
    services.repository.create_job(
        "integrity_sample", None, None, None, None, {"date": "2030-01-02", "sample_size": 20_000}, priority=1_000
    )

    succeeded = worker.run_once()

    assert succeeded is not None and succeeded["state"] == "succeeded"
    assert succeeded["message"] == "空闲完整性抽样完成"
    assert observed == [10_000]
    assert services.repository.get_settings()["last_integrity_sample_date"] == "2030-01-02"

    services.repository.update_settings({"max_retry_attempts": 0})
    services.repository.create_job(
        "integrity_sample", None, None, None, None, {"date": "2030-01-03", "sample_size": 1}, priority=1_000
    )
    failed = JobService(
        services.repository,
        services.artifacts,
        services.documents,
        integrity_runner=lambda _sample_size: {"valid": False, "failures": ["never-persisted"]},
    ).run_once()

    assert failed is not None and failed["state"] == "failed"
    assert failed["message"] == "本地处理失败"
    assert services.repository.get_settings()["last_integrity_sample_date"] == "2030-01-02"


def test_manual_full_integrity_verify_checks_every_catalogued_artifact(runtime_root: Path) -> None:
    services = create_app(runtime_root, acquire_lock=False).state.services
    services.imports.paste(PasteImportRequest(title="first", text="first artifact", rights="owned"))
    services.imports.paste(PasteImportRequest(title="second", text="second artifact", rights="owned"))

    result = services.transfers.verify_artifacts(full=True, sample_size=1)

    assert result["valid"] is True
    assert result["checked"] == 2


def test_external_userinfo_rejected_and_legacy_export_redacts(client: TestClient, runtime_root: Path) -> None:
    rejected = client.post("/api/v1/external/cards", json={"url": "https://user:secret@example.test/path", "title": "secret"})
    assert rejected.status_code == 422
    app = client.app
    with app.state.services.repository.connection() as connection:
        connection.execute(
            "INSERT INTO external_cards(id,card_type,url,title,author,notes,tags_json,created_at) VALUES(?,?,?,?,?,?,?,?)",
            ("legacy", "general", "https://user:legacy-secret@example.test/path", "legacy", None, None, "[]", "2030-01-01T00:00:00+00:00"),
        )
    assert "legacy-secret" not in json.dumps(client.get("/api/v1/external/cards").json())
    exported = client.post("/api/v1/exports", json={"confirmed": True}).json()
    with ZipFile(exported["archive_path"]) as archive:
        assert "legacy-secret" not in archive.read("records.json").decode()


def test_metadata_revisions_and_search_sort(client: TestClient) -> None:
    imported = paste(client, "Zulu title", "searchable")
    source_id = imported["source"]["id"]
    updated = client.put(f"/api/v1/sources/{source_id}/metadata", json={"title": "Alpha title"})
    assert updated.status_code == 200
    revisions = client.get(f"/api/v1/sources/{source_id}/metadata-revisions")
    assert revisions.status_code == 200
    assert len(revisions.json()) == 2
    assert revisions.json()[0]["snapshot"]["title"] == "Alpha title"
    parse(client)
    result = client.get("/api/v1/search", params={"q": "searchable", "sort": "title"})
    assert result.status_code == 200
    assert result.json()["sort"] == "title"
    assert client.get("/api/v1/search", params={"sort": "invalid"}).status_code == 422


def test_metadata_update_can_clear_nullable_fields_without_nulling_required_columns(client: TestClient) -> None:
    imported = client.post("/api/v1/imports/paste", json={
        "title": "带可清空元数据的来源",
        "text": "metadata clear",
        "rights": "owned",
        "author": "原作者",
        "notes": "原备注",
        "source_date": "2024-01-02",
    })
    assert imported.status_code == 201, imported.text
    source_id = imported.json()["source"]["id"]

    cleared = client.put(f"/api/v1/sources/{source_id}/metadata", json={
        "author": None,
        "notes": None,
        "source_date": None,
    })

    assert cleared.status_code == 200, cleared.text
    assert cleared.json()["author"] is None
    assert cleared.json()["notes"] is None
    assert cleared.json()["source_date"] is None
    assert cleared.json()["title"] == "带可清空元数据的来源"
    assert cleared.json()["language"] == "zh"
    revision = client.get(f"/api/v1/sources/{source_id}/metadata-revisions").json()[0]["snapshot"]
    assert revision["author"] is None
    assert revision["notes"] is None
    assert revision["source_date"] is None
    rejected = client.put(f"/api/v1/sources/{source_id}/metadata", json={"title": None})
    assert rejected.status_code == 422


def test_http_and_request_validation_errors_use_stable_envelopes(client: TestClient) -> None:
    missing = client.get("/api/v1/sources/missing")
    assert missing.status_code == 404
    assert missing.json()["detail"] == {"code": "http_404", "message": "来源不存在"}

    method_not_allowed = client.post("/api/v1/health")
    assert method_not_allowed.status_code == 405
    assert method_not_allowed.json()["detail"] == {"code": "http_405", "message": "请求方法不被允许"}

    invalid = client.put("/api/v1/settings", json={"job_lease_seconds": 1})
    assert invalid.status_code == 422
    assert invalid.json()["detail"] == {"code": "request_validation", "message": "请求字段无效"}


class _HangingProcess:
    def __init__(self) -> None:
        self.alive = True
        self.terminated = False

    def start(self) -> None:
        pass

    def is_alive(self) -> bool:
        return self.alive

    def terminate(self) -> None:
        self.terminated = True
        self.alive = False

    def join(self, *, timeout: float) -> None:
        assert timeout == 1


class _NoResultQueue:
    def __init__(self) -> None:
        self.closed = False

    def get(self, *, timeout: float):
        raise queue.Empty

    def get_nowait(self):
        raise queue.Empty

    def close(self) -> None:
        self.closed = True


def test_incomplete_pruning_record_is_reconciled_before_retention(runtime_root: Path) -> None:
    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services
    archive_name = "backup-interrupted.zip"
    archive_path = services.paths.backups / archive_name
    archive_path.write_bytes(b"interrupted backup")
    record = services.repository.create_backup_record(archive_name, "manifest", state="pruning")

    services.transfers._reconcile_incomplete_backup_records()

    reconciled = next(item for item in services.repository.list_backups() if item["id"] == record["id"])
    assert reconciled["state"] == "succeeded"
    assert archive_path.is_file()

    archive_path.unlink()
    services.repository.update_backup_state(record["id"], "pruning")
    services.transfers._reconcile_incomplete_backup_records()
    assert all(item["id"] != record["id"] for item in services.repository.list_backups())


def test_parser_timeout_circuit_breaker_terminates_without_waiting(runtime_root: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services
    process = _HangingProcess()
    result_queue = _NoResultQueue()

    class Context:
        def Queue(self, *, maxsize: int):
            assert maxsize == 1
            return result_queue

        def Process(self, **kwargs):
            return process

    times = iter((0.0, 1.0))
    monkeypatch.setattr("app.services.jobs.multiprocessing.get_context", lambda mode: Context())
    monkeypatch.setattr("app.services.jobs.time.monotonic", lambda: next(times))

    worker = JobService(services.repository, services.artifacts, services.documents)
    with pytest.raises(ParserCircuitBreaker, match="解析超时断路器已触发"):
        worker._run_parser_with_circuit_breakers(
            runtime_root / "artifact", "fixture.md", "text/markdown", 1.0, 60.0, lambda: False, lambda: None
        )
    assert process.terminated is True
    assert result_queue.closed is True


def test_retention_delete_failure_keeps_archive_and_success_record(runtime_root: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services
    archive_name = "backup-retention-failure.zip"
    archive_path = services.paths.backups / archive_name
    archive_path.write_bytes(b"retained")
    record = services.repository.create_backup_record(archive_name, "manifest", state="succeeded")
    original_unlink = Path.unlink

    def deny_archive_unlink(path: Path, *args, **kwargs):
        if path == archive_path:
            raise OSError("injected")
        return original_unlink(path, *args, **kwargs)

    monkeypatch.setattr(Path, "unlink", deny_archive_unlink)
    services.transfers._prune_backups(0)

    retained = next(item for item in services.repository.list_backups() if item["id"] == record["id"])
    assert retained["state"] == "succeeded"
    assert archive_path.is_file()


def test_failed_backup_keeps_discarding_record_when_archive_cleanup_fails(runtime_root: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    app = create_app(runtime_root, acquire_lock=False)
    services = app.state.services
    original_unlink = Path.unlink

    def deny_backup_unlink(path: Path, *args, **kwargs):
        if path.parent == services.paths.backups:
            raise OSError("injected")
        return original_unlink(path, *args, **kwargs)

    monkeypatch.setattr(Path, "unlink", deny_backup_unlink)
    monkeypatch.setattr(services.transfers, "_prune_backups", lambda *args, **kwargs: (_ for _ in ()).throw(RuntimeError("injected")))
    with pytest.raises(RuntimeError, match="injected"):
        services.transfers.create_backup()

    records = services.repository.list_backups()
    assert len(records) == 1
    assert records[0]["state"] == "discarding"
    assert (services.paths.backups / records[0]["archive_name"]).is_file()
